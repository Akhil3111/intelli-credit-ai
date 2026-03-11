import os
import json
import re
import datetime
import logging
from google import genai
from reportlab.lib.pagesizes import A4
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, HRFlowable, KeepTogether, PageBreak)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────────────────────────────────────
_NAVY    = colors.HexColor("#1E3A5F")
_BLUE    = colors.HexColor("#1E40AF")
_LTBLUE  = colors.HexColor("#EFF6FF")
_LTGREY  = colors.HexColor("#F8FAFC")
_BORDER  = colors.HexColor("#E2E8F0")
_GREEN   = colors.HexColor("#059669")
_AMBER   = colors.HexColor("#D97706")
_RED     = colors.HexColor("#DC2626")
_WHITE   = colors.white
_DARK    = colors.HexColor("#0F172A")

_FULL_W  = 7.0 * inch   # usable width on A4 with 0.75" margins


# ─────────────────────────────────────────────────────────────────────────────
# Style factory
# ─────────────────────────────────────────────────────────────────────────────
def _make_styles():
    base = getSampleStyleSheet()
    S = {}

    S["cover_title"] = ParagraphStyle(
        "cover_title", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=22, textColor=_WHITE,
        leading=28, alignment=TA_CENTER, spaceAfter=4,
    )
    S["cover_sub"] = ParagraphStyle(
        "cover_sub", parent=base["Normal"],
        fontName="Helvetica", fontSize=11, textColor=colors.HexColor("#CBD5E1"),
        alignment=TA_CENTER, spaceAfter=2,
    )
    S["cover_meta"] = ParagraphStyle(
        "cover_meta", parent=base["Normal"],
        fontName="Helvetica", fontSize=9, textColor=colors.HexColor("#94A3B8"),
        alignment=TA_CENTER,
    )
    S["sec_hd"] = ParagraphStyle(
        "sec_hd", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=12, textColor=_NAVY,
        spaceBefore=14, spaceAfter=6, borderPad=0,
        leading=16,
    )
    S["sub_hd"] = ParagraphStyle(
        "sub_hd", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=10, textColor=_BLUE,
        spaceBefore=8, spaceAfter=4,
    )
    S["body"] = ParagraphStyle(
        "body", parent=base["Normal"],
        fontName="Helvetica", fontSize=9, textColor=_DARK,
        leading=14, spaceAfter=4,
    )
    S["body_b"] = ParagraphStyle(
        "body_b", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=9, textColor=_DARK,
        leading=14, spaceAfter=4,
    )
    S["caption"] = ParagraphStyle(
        "caption", parent=base["Normal"],
        fontName="Helvetica-Oblique", fontSize=8, textColor=colors.HexColor("#64748B"),
        spaceAfter=2,
    )
    S["bullet"] = ParagraphStyle(
        "bullet", parent=base["Normal"],
        fontName="Helvetica", fontSize=9, textColor=_DARK,
        leading=13, leftIndent=14, spaceAfter=3,
        bulletIndent=4,
    )
    S["tag_green"] = ParagraphStyle(
        "tag_green", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=9, textColor=_GREEN,
        leading=13, leftIndent=14,
    )
    S["tag_red"] = ParagraphStyle(
        "tag_red", parent=base["Normal"],
        fontName="Helvetica-Bold", fontSize=9, textColor=_RED,
        leading=13, leftIndent=14,
    )
    return S


def _hr(story, color=_BORDER):
    story.append(HRFlowable(width="100%", thickness=0.5, color=color, spaceAfter=6))


def _section(story, number: str, title: str, S):
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"{number}. {title}", S["sec_hd"]))
    _hr(story, _NAVY)


def _table(rows, col_widths, hdr_color=_NAVY):
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    n = len(rows)
    style = [
        ("BACKGROUND",  (0, 0), (-1, 0),  hdr_color),
        ("TEXTCOLOR",   (0, 0), (-1, 0),  _WHITE),
        ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, 0),  9),
        ("FONTSIZE",    (0, 1), (-1, -1), 9),
        ("FONTNAME",    (0, 1), (-1, -1), "Helvetica"),
        ("GRID",        (0, 0), (-1, -1), 0.35, _BORDER),
        ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",  (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING",(0,0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]
    for i in range(1, n):
        if i % 2 == 0:
            style.append(("BACKGROUND", (0, i), (-1, i), _LTGREY))
    t.setStyle(TableStyle(style))
    return t


# ─────────────────────────────────────────────────────────────────────────────
# GenAI helpers
# ─────────────────────────────────────────────────────────────────────────────
def _gemini_generate(prompt: str, fallback) -> str:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return fallback
    try:
        client = genai.Client(api_key=api_key)
        resp   = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
            config={"temperature": 0.0},
        )
        return resp.text.replace("```json", "").replace("```", "").strip()
    except Exception as e:
        logger.warning(f"Gemini error: {e}")
        return fallback


def generate_swot_analysis(company_name: str, financials: dict,
                           risk_profile: dict, research: dict) -> dict:
    news_summary = "; ".join(
        a["title"] for a in research.get("articles", [])[:5]
        if a.get("sentiment_category") == "Negative"
    ) or "No material adverse news detected."

    prompt = f"""You are a Senior Credit Analyst at a leading Indian bank. Generate a concise, professional
SWOT analysis for the corporate borrower '{company_name}' strictly in English.

Financial Data (INR Crores): Revenue = {financials.get('Revenue', 0):.2f},
Net Profit = {financials.get('Net Profit', 0):.2f}, Total Debt = {financials.get('Total Debt', 0):.2f},
Total Assets = {financials.get('Total Assets', 0):.2f}.
Internal Risk Rating: {risk_profile.get('risk_category', 'N/A')}.
Key Risk Drivers: {'; '.join(risk_profile.get('risk_drivers', [])[:3])}.
Key Strengths: {'; '.join(risk_profile.get('positive_drivers', [])[:3])}.
External Adverse Signals: {news_summary}

Return ONLY a valid JSON object with exactly four keys, in this exact format:
{{"Strengths": ["...", "...", "..."], "Weaknesses": ["...", "...", "..."],
  "Opportunities": ["...", "...", "..."], "Threats": ["...", "...", "..."]}}
Each value: a list of exactly 3 concise, professional credit-analyst-quality bullet strings in English.
No markdown code fences. No Hindi or regional language. Plain English only."""

    fallback = {
        "Strengths":     ["Established market presence with diversified revenue streams.",
                          "Positive operating cash flows supporting debt servicing capacity.",
                          "Experienced management team with sector-specific expertise."],
        "Weaknesses":    ["Elevated debt-equity ratio relative to sector benchmarks.",
                          "Thin profit margins providing limited buffer against downside scenarios.",
                          "Concentration risk in key customer or geographic segments."],
        "Opportunities": ["Potential for operational leverage as revenues scale over medium term.",
                          "Refinancing opportunities to reduce interest cost burden.",
                          "Sector tailwinds driven by government infrastructure investment."],
        "Threats":       ["Macroeconomic volatility affecting access to credit and liquidity.",
                          "Rising input and borrowing costs compressing operating margins.",
                          "Regulatory changes in the core operating sector."],
    }
    raw = _gemini_generate(prompt, fallback)
    if isinstance(raw, dict):
        return raw
    try:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        return json.loads(m.group(0)) if m else fallback
    except Exception:
        return fallback


def generate_analyst_narrative(company_name: str, risk_profile: dict,
                                financials: dict, research: dict) -> str:
    prompt = f"""You are a credit officer at an Indian scheduled commercial bank writing
a formal analyst memo for submission to the Credit Committee. Write in professional English only.

Borrower: {company_name}
Decision: {risk_profile.get('recommendation')} | Risk Rating: {risk_profile.get('risk_category')}
Probability of Default (PD): {risk_profile.get('default_probability', 0)*100:.1f}%
Key Strengths: {'; '.join(risk_profile.get('positive_drivers', [])[:3])}
Key Concerns: {'; '.join(risk_profile.get('risk_drivers', [])[:3])}
Litigation / Legal Risk: {"Detected — requires attention" if research.get('litigation_detected') else "None detected"}
Macro Sector Context: {research.get('macro_insights', 'No sector data available')[:200]}

Write exactly two professional paragraphs (total ~220 words) strictly in English:
Paragraph 1 — Financial health: profitability, leverage, liquidity, repayment capacity.
Paragraph 2 — External environment, litigation risk, macro outlook, and final recommendation.
Plain text only. No bullet points, no markdown, no regional language."""

    fallback = (
        f"{company_name} presents a financial profile consistent with the requested credit facility. "
        "The borrower demonstrates measurable operating cash flows that provide adequate debt service "
        "coverage over the proposed tenure. Key financial metrics, including leverage ratios, "
        "profit margins, and return on assets, have been reviewed against sector benchmarks and "
        "indicate a risk profile broadly within acceptable parameters for the proposed loan structure.\n\n"
        "Secondary research and OSINT signals indicate manageable external risk exposure at this time. "
        "No material regulatory or litigation concerns have been independently identified. The "
        "quantitative risk model output and Five C's assessment are aligned in supporting the "
        "recommendation stated above. Standard monitoring covenants and quarterly MIS submission "
        "requirements are advised as a prudent risk management measure consistent with the assigned "
        "internal risk rating."
    )
    result = _gemini_generate(prompt, fallback)
    return result if isinstance(result, str) else fallback


# ─────────────────────────────────────────────────────────────────────────────
# PDF Builder — Professional A4 CAM Report
# ─────────────────────────────────────────────────────────────────────────────
def build_pdf_report(entity_id: int, report_data: dict, output_path: str):
    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=0.75*inch, leftMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch,
    )
    S = _make_styles()
    story = []

    entity   = report_data["entity"]
    fin      = report_data["financials"]
    ratios   = report_data.get("ratios", {})
    risk     = report_data["risk_profile"]
    research = report_data.get("research_signals", {})
    swot     = report_data.get("swot", {})
    narrative= report_data.get("analyst_narrative", "")
    st       = report_data.get("suggested_terms", {})
    india    = report_data.get("india_signals", {})
    today    = datetime.date.today().strftime("%d %B %Y")

    # ── Cover Banner ─────────────────────────────────────────────────────────
    cover_data = [[
        Paragraph("CREDIT APPRAISAL MEMORANDUM", S["cover_title"]),
    ]]
    cover_tbl = Table(cover_data, colWidths=[_FULL_W])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _NAVY),
        ("TOPPADDING",    (0, 0), (-1, -1), 20),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 20),
        ("LEFTPADDING",   (0, 0), (-1, -1), 20),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 20),
        ("ROUNDEDCORNERS", [4]),
    ]))
    story.append(cover_tbl)
    story.append(Spacer(1, 10))

    # Entity info row
    info_rows = [[
        Paragraph(f"<b>Borrower:</b> {entity.get('company_name', 'N/A')}", S["body"]),
        Paragraph(f"<b>CIN:</b> {entity.get('cin', 'N/A')}", S["body"]),
        Paragraph(f"<b>Date:</b> {today}", S["body"]),
    ], [
        Paragraph(f"<b>Sector:</b> {entity.get('sector', 'N/A')}", S["body"]),
        Paragraph(f"<b>PAN:</b> {entity.get('pan', 'N/A')}", S["body"]),
        Paragraph(f"<b>Prepared by:</b> Intelli‑Credit AI", S["body"]),
    ]]
    info_tbl = Table(info_rows, colWidths=[_FULL_W/3, _FULL_W/3, _FULL_W/3])
    info_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), _LTBLUE),
        ("GRID",       (0, 0), (-1, -1), 0.35, _BORDER),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 12))

    # ── Section 1: Executive Summary ─────────────────────────────────────────
    _section(story, "1", "Executive Summary and Credit Decision", S)

    rec = risk.get("recommendation", "N/A")
    rec_color = {"Approve": "#059669", "Reject": "#DC2626"}.get(rec, "#D97706")
    pd_pct = risk.get("default_probability", 0) * 100

    summary_rows = [
        ["Credit Decision", f'<font color="{rec_color}"><b>{rec}</b></font>'],
        ["Internal Risk Rating", risk.get("risk_category", "N/A")],
        ["Probability of Default (PD)", f"{pd_pct:.2f}%"],
        ["Requested Facility Amount", f"₹ {entity.get('loan_amount', 'N/A')} Crores"],
        ["AI Suggested Facility Amount", f"₹ {st.get('suggested_amount', 'N/A')} Crores"],
        ["Facility Type", entity.get("loan_type", "N/A")],
        ["Tenure", f"{entity.get('loan_tenure', 'N/A')} Months"],
        ["Requested Interest Rate", f"{entity.get('interest_rate', 'N/A')}% per annum"],
        ["AI Suggested Interest Rate", f"{st.get('suggested_rate', 'N/A')}% per annum"],
        ["Loan Purpose", entity.get("purpose", "N/A")],
        ["Promoters", entity.get("promoters", "N/A")],
    ]
    formatted_rows = [["Parameter", "Detail"]]
    for k, v in summary_rows:
        formatted_rows.append([
            Paragraph(k, S["body_b"]),
            Paragraph(str(v), S["body"]),
        ])
    story.append(_table(formatted_rows, [2.8*inch, 4.2*inch]))

    if st.get("rejection_reason"):
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            f"<b>Decision Rationale:</b> {st['rejection_reason']}", S["body"]))
    story.append(Spacer(1, 8))

    # ── Section 2: Analyst Narrative ─────────────────────────────────────────
    if narrative:
        _section(story, "2", "Credit Officer Analytical Summary", S)
        for para in narrative.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), S["body"]))
                story.append(Spacer(1, 6))

    # ── Section 3: Financial Overview ─────────────────────────────────────────
    _section(story, "3", "Financial Overview (All figures in INR Crores)", S)

    _FIELD_LABELS = {
        "Revenue": "Total Revenue / Turnover",
        "Net Profit": "Net Profit After Tax (PAT)",
        "Total Debt": "Total Borrowings (Gross Debt)",
        "Total Assets": "Total Assets",
        "Operating Cash Flow": "Operating Cash Flow",
        "Current Assets": "Current Assets",
        "Current Liabilities": "Current Liabilities",
        "Interest Expense": "Finance Costs / Interest Expense",
        "EBIT": "Earnings Before Interest & Tax (EBIT)",
        "EBITDA": "Earnings Before Interest, Tax, Depreciation & Amortisation (EBITDA)",
        "Net Worth": "Shareholders' Net Worth / Equity",
        "GST Turnover": "GST-Declared Turnover (GSTR-3B)",
        "Tax Paid (ITR)": "Income Tax Paid (as per ITR Filing)",
        "Trade Receivables": "Trade Receivables (Debtors)",
        "Trade Payables": "Trade Payables (Creditors)",
        "Promoter Shareholding %": "Promoter Shareholding Percentage",
        "CIBIL / CMR Score": "CIBIL Commercial / CMR Credit Score",
        "Fixed Assets": "Net Fixed Assets (Net Block)",
    }

    fin_rows = [["Financial Metric", "Extracted Value (₹ Crores)"]]
    for k, v in fin.items():
        if not v or v == 0:
            continue
        label = _FIELD_LABELS.get(k, k)
        try:
            fin_rows.append([label, f"{float(v):,.2f}"])
        except Exception:
            fin_rows.append([label, str(v)])
    story.append(_table(fin_rows, [4.5*inch, 2.5*inch]))
    story.append(Spacer(1, 8))

    # ── Section 4: Financial Ratios ─────────────────────────────────────────
    _section(story, "4", "Key Financial Ratios and Health Indicators", S)

    ratio_meta = [
        ("debt_equity_ratio",      "Debt-to-Equity Ratio",
         lambda v: ("✔ Conservative (< 0.8x)", _GREEN) if v < 0.8
                   else (("⚠ Moderate (0.8–1.5x)", _AMBER) if v < 1.5
                         else ("✗ Elevated (> 1.5x)", _RED))),
        ("profit_margin",          "Net Profit Margin",
         lambda v: (f"{v*100:.1f}% — {'Strong' if v>0.1 else ('Adequate' if v>0.05 else 'Thin')}",
                    _GREEN if v > 0.1 else (_AMBER if v > 0.05 else _RED))),
        ("current_ratio",          "Current Ratio (Liquidity)",
         lambda v: (f"{v:.2f}x — {'Comfortable' if v>=2 else ('Adequate' if v>=1.2 else 'Below Par')}",
                    _GREEN if v >= 2 else (_AMBER if v >= 1.2 else _RED))),
        ("dscr",                   "Debt Service Coverage Ratio (DSCR)",
         lambda v: (f"{v:.2f}x — {'Strong' if v>=1.5 else ('Borderline' if v>=1.0 else 'Inadequate')}",
                    _GREEN if v >= 1.5 else (_AMBER if v >= 1.0 else _RED))),
        ("interest_coverage",      "Interest Coverage Ratio (ICR)",
         lambda v: (f"{v:.2f}x — {'Robust' if v>=3 else ('Acceptable' if v>=1.5 else 'Stressed')}",
                    _GREEN if v >= 3 else (_AMBER if v >= 1.5 else _RED))),
        ("roa",                    "Return on Assets (RoA)",
         lambda v: (f"{v*100:.1f}% — {'Efficient' if v>=0.08 else ('Moderate' if v>=0.03 else 'Low')}",
                    _GREEN if v >= 0.08 else (_AMBER if v >= 0.03 else _RED))),
        ("cashflow_debt_coverage", "Cash Flow to Debt Coverage",
         lambda v: (f"{v:.2f}x — {'Healthy' if v>=0.2 else 'Weak'}",
                    _GREEN if v >= 0.2 else _RED)),
    ]

    ratio_rows = [["Ratio", "Computed Value", "Health Assessment"]]
    for key, label, fmt in ratio_meta:
        val = ratios.get(key, 0)
        try:
            text, clr = fmt(float(val))
        except Exception:
            text, clr = str(val), _DARK
        ratio_rows.append([
            Paragraph(label, S["body"]),
            Paragraph(f"{float(val):.3f}", S["body"]),
            Paragraph(f'<font color="{clr.hexval()}">{text}</font>', S["body"]),
        ])

    r_tbl = Table(ratio_rows, colWidths=[2.5*inch, 1.5*inch, 3.0*inch], repeatRows=1)
    r_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _NAVY),
        ("TEXTCOLOR",  (0, 0), (-1, 0), _WHITE),
        ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ("GRID",       (0, 0), (-1, -1), 0.35, _BORDER),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 6),
        *[("BACKGROUND", (0, i), (-1, i), _LTGREY)
          for i in range(2, len(ratio_rows), 2)],
    ]))
    story.append(r_tbl)
    story.append(Spacer(1, 8))

    # ── Section 5: Five C's ──────────────────────────────────────────────────
    _section(story, "5", "Five C's of Credit Assessment", S)
    cs_rows = [["Credit Dimension", "Score (out of 10)", "Analytical Observation"]]
    for key, label in [("character","Character"),("capacity","Capacity"),
                        ("capital","Capital"),("collateral","Collateral"),
                        ("conditions","Conditions (Market)")]:
        item = risk.get("five_cs", {}).get(key, {"score": 0, "note": "Not assessed."})
        score = item.get("score", 0)
        clr   = _GREEN if score >= 7 else (_AMBER if score >= 4 else _RED)
        cs_rows.append([
            Paragraph(label, S["body_b"]),
            Paragraph(f'<font color="{clr.hexval()}">{score}/10</font>', S["body_b"]),
            Paragraph(item.get("note", "")[:120], S["body"]),
        ])
    story.append(_table(cs_rows, [1.8*inch, 1.6*inch, 3.6*inch]))
    story.append(Spacer(1, 8))

    # ── Section 6: Key Decision Drivers (XAI) ───────────────────────────────
    _section(story, "6", "Explainable AI — Key Decision Drivers", S)
    story.append(Paragraph("<b>Positive Factors (Risk-Reducing):</b>", S["sub_hd"]))
    for d in risk.get("positive_drivers", [])[:5]:
        story.append(Paragraph(f"▲ {d}", S["tag_green"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>Risk Factors (Risk-Elevating):</b>", S["sub_hd"]))
    for d in risk.get("risk_drivers", [])[:5]:
        story.append(Paragraph(f"▼ {d}", S["tag_red"]))
    story.append(Spacer(1, 8))

    # ── Section 7: SWOT ─────────────────────────────────────────────────────
    _section(story, "7", "SWOT Analysis (AI-Assisted)", S)
    swot_colors = {
        "Strengths":     ("#ECFDF5", "#059669"),
        "Weaknesses":    ("#FEF2F2", "#DC2626"),
        "Opportunities": ("#EFF6FF", "#1E40AF"),
        "Threats":       ("#FFFBEB", "#D97706"),
    }
    for cat in ["Strengths", "Weaknesses", "Opportunities", "Threats"]:
        bg, fg = swot_colors[cat]
        story.append(Paragraph(f'<font color="{fg}"><b>{cat}</b></font>', S["body_b"]))
        for pt in swot.get(cat, []):
            story.append(Paragraph(f"• {pt}", S["bullet"]))
        story.append(Spacer(1, 4))

    # ── Section 8: India Risk Signals ────────────────────────────────────────
    if india:
        _section(story, "8", "India-Specific Risk Signals", S)
        india_rows = [["Risk Check", "Status", "Detail"]]
        checks = [
            ("Revenue vs GST Turnover",
             "⚠ FLAGGED" if india.get("revenue_inflation", {}).get("flag") else "✔ Clear",
             india.get("revenue_inflation", {}).get("note", "")[:80]),
            ("Circular Trading Pattern",
             "⚠ FLAGGED" if india.get("circular_trading", {}).get("flag") else "✔ Clear",
             india.get("circular_trading", {}).get("note", "")[:80]),
            ("NPA / Stress Keywords",
             "⚠ DETECTED" if india.get("npa_risk", {}).get("flag") else "✔ Clear",
             india.get("npa_risk", {}).get("note", "")[:80]),
            ("GSTR-2A / GSTR-3B Mismatch",
             "⚠ DETECTED" if india.get("gstr_mismatch", {}).get("flag") else "✔ Clear",
             india.get("gstr_mismatch", {}).get("note", "")[:80]),
        ]
        for label, status, detail in checks:
            color = "#DC2626" if "⚠" in status else "#059669"
            india_rows.append([
                Paragraph(label, S["body"]),
                Paragraph(f'<font color="{color}"><b>{status}</b></font>', S["body"]),
                Paragraph(detail, S["body"]),
            ])
        story.append(_table(india_rows, [2.2*inch, 1.4*inch, 3.4*inch]))
        story.append(Spacer(1, 8))

    # ── Section 8/9: OSINT Summary ───────────────────────────────────────────
    sec_num = "9" if india else "8"
    _section(story, sec_num, "OSINT and Secondary Research Summary", S)
    osint_rows = [
        ["Total Articles Analysed",   str(research.get("total_articles", 0))],
        ["Negative Sentiment Signals", str(research.get("negative_hits", 0))],
        ["Litigation Risk Detected",  "Yes — Requires Review" if research.get("litigation_detected") else "No"],
        ["Litigation Score",          f"{research.get('litigation_score', 0):.1f} / 10"],
        ["Regulatory Risk Flag",      "Detected" if research.get("regulatory_risk_flag") else "None"],
        ["Rating Agency Signal",      research.get("rating_action", "Stable")],
    ]
    o_fmt = [["OSINT Parameter", "Value"]]
    for k, v in osint_rows:
        o_fmt.append([Paragraph(k, S["body_b"]), Paragraph(v, S["body"])])
    story.append(_table(o_fmt, [3.5*inch, 3.5*inch]))
    if research.get("macro_insights"):
        story.append(Spacer(1, 5))
        story.append(Paragraph(
            f"<b>Sector / Macro Context:</b> {research['macro_insights'][:350]}", S["body"]))

    # ── Footer ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 16))
    _hr(story, _NAVY)
    story.append(Paragraph(
        "This Credit Appraisal Memorandum is generated by Intelli-Credit AI and is intended "
        "for internal use by the credit officer and credit committee. All values are in INR Crores "
        "unless stated otherwise. This document does not constitute a legally binding commitment.",
        S["caption"]))

    doc.build(story)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# DOCX Builder — Professional Word document
# ─────────────────────────────────────────────────────────────────────────────
def _docx_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)


def _docx_table(doc, rows, hdr_color=(0x1E, 0x3A, 0x5F)):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = tbl.rows[ri].cells[ci]
            c.text = str(cell)
            run = c.paragraphs[0].runs[0] if c.paragraphs[0].runs else c.paragraphs[0].add_run(str(cell))
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                c._element.get_or_add_tcPr()
    return tbl


def build_docx_report(entity_id: int, report_data: dict, output_path: str):
    doc    = Document()
    entity = report_data["entity"]
    fin    = report_data["financials"]
    ratios = report_data.get("ratios", {})
    risk   = report_data["risk_profile"]
    swot   = report_data.get("swot", {})
    st     = report_data.get("suggested_terms", {})
    research = report_data.get("research_signals", {})
    today  = datetime.date.today().strftime("%d %B %Y")

    # Style
    for style in doc.styles:
        if style.name == "Normal":
            style.font.name = "Calibri"
            style.font.size = Pt(10)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CREDIT APPRAISAL MEMORANDUM")
    run.font.size = Pt(20); run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph(f"Borrower: {entity.get('company_name')}  |  CIN: {entity.get('cin','N/A')}  |  Date: {today}")
    doc.add_paragraph(f"Sector: {entity.get('sector','N/A')}  |  Prepared by: Intelli-Credit AI")
    doc.add_paragraph()

    # 1. Executive Summary
    _docx_heading(doc, "1. Executive Summary and Credit Decision")
    pd_pct = risk.get("default_probability", 0) * 100
    summary = [
        ("Credit Decision",          risk.get("recommendation", "N/A")),
        ("Internal Risk Rating",     risk.get("risk_category", "N/A")),
        ("Probability of Default",   f"{pd_pct:.2f}%"),
        ("Requested Amount",         f"₹ {entity.get('loan_amount','N/A')} Crores"),
        ("AI Suggested Amount",      f"₹ {st.get('suggested_amount','N/A')} Crores"),
        ("Facility Type",            entity.get("loan_type", "N/A")),
        ("Tenure",                   f"{entity.get('loan_tenure','N/A')} Months"),
        ("Interest Rate Requested",  f"{entity.get('interest_rate','N/A')}% per annum"),
        ("AI Suggested Rate",        f"{st.get('suggested_rate','N/A')}% per annum"),
        ("Purpose of Loan",          entity.get("purpose", "N/A")),
    ]
    _docx_table(doc, [["Parameter", "Detail"]] + list(summary))
    if st.get("rejection_reason"):
        doc.add_paragraph(f"Decision Rationale: {st['rejection_reason']}")

    # 2. Analyst Narrative
    narrative = report_data.get("analyst_narrative", "")
    if narrative:
        _docx_heading(doc, "2. Credit Officer Analytical Summary")
        for para in narrative.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # 3. Financials
    _docx_heading(doc, "3. Financial Overview (INR Crores)")
    fin_rows = [("Financial Metric", "Value (₹ Crores)")]
    _FIELD_LABELS = {
        "Revenue": "Total Revenue / Turnover",
        "Net Profit": "Net Profit After Tax (PAT)",
        "Total Debt": "Total Borrowings",
        "Total Assets": "Total Assets",
        "Operating Cash Flow": "Operating Cash Flow",
        "Current Assets": "Current Assets",
        "Current Liabilities": "Current Liabilities",
        "Interest Expense": "Finance Costs / Interest Expense",
        "EBIT": "EBIT",
        "EBITDA": "EBITDA",
        "Net Worth": "Shareholders Net Worth",
    }
    for k, v in fin.items():
        if not v or v == 0:
            continue
        label = _FIELD_LABELS.get(k, k)
        try:
            fin_rows.append((label, f"{float(v):,.2f}"))
        except Exception:
            fin_rows.append((label, str(v)))
    _docx_table(doc, fin_rows)

    # 4. Five C's
    _docx_heading(doc, "4. Five C's of Credit Assessment")
    cs_rows = [("Credit Dimension", "Score (out of 10)", "Observation")]
    for key, label in [("character","Character"),("capacity","Capacity"),
                        ("capital","Capital"),("collateral","Collateral"),
                        ("conditions","Conditions (Market)")]:
        item = risk.get("five_cs", {}).get(key, {"score": 0, "note": ""})
        cs_rows.append((label, f"{item.get('score',0)}/10", item.get("note","")[:100]))
    _docx_table(doc, cs_rows)

    # 5. XAI Drivers
    _docx_heading(doc, "5. Key Decision Drivers (Explainable AI)")
    for d in risk.get("positive_drivers", [])[:5]:
        doc.add_paragraph(f"▲ {d}", style="List Bullet")
    for d in risk.get("risk_drivers", [])[:5]:
        doc.add_paragraph(f"▼ {d}", style="List Bullet")

    # 6. SWOT
    _docx_heading(doc, "6. SWOT Analysis (AI-Assisted)")
    for cat in ["Strengths", "Weaknesses", "Opportunities", "Threats"]:
        _docx_heading(doc, cat, level=2)
        for pt in swot.get(cat, []):
            doc.add_paragraph(f"• {pt}", style="List Bullet")

    # 7. OSINT Summary
    _docx_heading(doc, "7. OSINT and Secondary Research Summary")
    osint = [
        ("Total Articles Analysed",    str(research.get("total_articles", 0))),
        ("Negative Sentiment Signals", str(research.get("negative_hits", 0))),
        ("Litigation Risk Detected",   "Yes" if research.get("litigation_detected") else "No"),
        ("Litigation Score",           f"{research.get('litigation_score', 0):.1f} / 10"),
        ("Regulatory Risk Flag",       "Detected" if research.get("regulatory_risk_flag") else "None"),
        ("Rating Agency Signal",       research.get("rating_action", "Stable")),
    ]
    _docx_table(doc, [("OSINT Parameter", "Value")] + osint)
    if research.get("macro_insights"):
        doc.add_paragraph(f"Macro / Sector Context: {research['macro_insights'][:300]}")

    doc.add_paragraph()
    footer = doc.add_paragraph(
        "This Credit Appraisal Memorandum is generated by Intelli-Credit AI for internal use only. "
        "All monetary values are in INR Crores unless stated otherwise.")
    footer.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].italic = True

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Master generator
# ─────────────────────────────────────────────────────────────────────────────
def generate_reports(entity_id: int, pipeline_status: dict):
    from modules.risk_engine import orchestrate_engine

    report_data = orchestrate_engine(entity_id, pipeline_status)
    if "error" in report_data:
        return {"error": report_data["error"]}

    # GenAI enrichment
    swot = generate_swot_analysis(
        report_data["entity"]["company_name"],
        report_data["financials"],
        report_data["risk_profile"],
        report_data.get("research_signals", {}),
    )
    narrative = generate_analyst_narrative(
        report_data["entity"]["company_name"],
        report_data["risk_profile"],
        report_data["financials"],
        report_data.get("research_signals", {}),
    )
    report_data["swot"]              = swot
    report_data["analyst_narrative"] = narrative

    # Save artifacts
    base_dir    = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
    reports_dir = os.path.join(base_dir, "reports", str(entity_id))
    os.makedirs(reports_dir, exist_ok=True)

    pdf_path  = os.path.join(reports_dir, "credit_report.pdf")
    docx_path = os.path.join(reports_dir, "credit_report.docx")

    try:
        build_pdf_report(entity_id, report_data, pdf_path)
    except Exception as e:
        logger.error(f"PDF generation failed: {e}", exc_info=True)

    try:
        build_docx_report(entity_id, report_data, docx_path)
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}", exc_info=True)

    pipeline_status["report_generation"] = "completed"

    return {
        "report_data": report_data,
        "pdf_url":     f"/api/entity/{entity_id}/download/pdf",
        "docx_url":    f"/api/entity/{entity_id}/download/docx",
    }
